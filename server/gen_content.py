import os
import random
import string
import zipfile
from pathlib import Path
import struct
from PIL import Image
import numpy as np
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from pptx import Presentation
import io
import logging
import concurrent.futures
import tempfile
from tqdm import tqdm

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

def get_size_string(size_kb):
    """Convert size in KB to human readable string."""
    if size_kb >= 1024 * 1024:  # >= 1GB
        return f"{size_kb/(1024*1024):.0f}GB"
    elif size_kb >= 1024:  # >= 1MB
        return f"{size_kb/1024:.0f}MB"
    else:
        return f"{size_kb:.0f}KB"

def generate_elf_binary(size_mb):
    """Generate a non-malicious ELF binary of approximately specified size."""
    elf_header = bytes.fromhex('7f454c46020101000000000000000000')
    padding = os.urandom(int(size_mb * 1024 * 1024) - len(elf_header))
    return elf_header + padding

def generate_document(size_mb, doc_type):
    """Generate a valid document of specified type and approximate size."""
    target_size = size_mb * 1024 * 1024
    
    if size_mb > 50 and doc_type in ['docx', 'pdf', 'pptx']:
        return None
    if doc_type == 'xlsx' and size_mb > 10:
        return None

    if doc_type == 'docx':
        doc = Document()
        pbar = tqdm(total=target_size, desc=f"Generating {size_mb:.1f}MB DOCX", unit='B', unit_scale=True)
        
        # Pre-generate some content blocks
        text_block = ''.join(random.choices(string.ascii_letters, k=1024))
        img_data = io.BytesIO()
        img_size = int(min(512, np.sqrt(target_size/10)))  # Smaller images, more of them
        img = Image.fromarray(np.random.randint(0, 255, (img_size, img_size, 3), dtype=np.uint8))
        img.save(img_data, format='PNG')
        img_data.seek(0)
        
        buffer = io.BytesIO()
        current_size = 0
        
        while current_size < target_size:
            if random.random() < 0.7:  # 70% images
                doc.add_picture(img_data)
                img_data.seek(0)
            else:
                doc.add_paragraph(text_block)
            
            buffer.seek(0)
            doc.save(buffer)
            new_size = buffer.tell()
            pbar.update(new_size - current_size)
            current_size = new_size
            if current_size >= target_size:
                break
        
        pbar.close()
        return buffer.getvalue()

    elif doc_type == 'pdf':
        buffer = io.BytesIO()
        pbar = tqdm(total=target_size, desc=f"Generating {size_mb:.1f}MB PDF", unit='B', unit_scale=True)
        
        c = canvas.Canvas(buffer)
        current_size = 0
        text_block = ''.join(random.choices(string.ascii_letters, k=1024))
        page = 1
        
        while current_size < target_size and page < 1000:  # Page limit for safety
            c.drawString(100, 100, text_block)
            c.showPage()
            page += 1
            
            if page % 10 == 0:  # Check size periodically
                c.save()
                new_size = buffer.tell()
                pbar.update(new_size - current_size)
                current_size = new_size
                if current_size >= target_size:
                    break
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer)
        
        pbar.close()
        return buffer.getvalue()

def generate_image(size_mb):
    """Generate an image of specified size."""
    dim = int(np.sqrt((size_mb * 1024 * 1024) / 3))  # 3 channels (RGB)
    img = Image.fromarray(np.random.randint(0, 255, (dim, dim, 3), dtype=np.uint8))
    buffer = io.BytesIO()
    img.save(buffer, format='JPEG', quality=95)
    return buffer.getvalue()

def create_size_variants(base_size_kb=20):
    """Generate list of sizes from 20KB to 300MB, doubling each time."""
    sizes_kb = []
    current_size = base_size_kb
    while current_size <= 520 * 1024:  # 300MB limit
        sizes_kb.append(current_size)
        current_size *= 2
    return sizes_kb

def generate_size_variant(size_kb, output_dir):
    size_mb = size_kb / 1024
    size_str = get_size_string(size_kb)
    config_lines = []
    
    try:
        # Generate binary file with progress bar
        binary_name = f"binary_{size_str}.bin"
        with tqdm(total=size_mb*1024*1024, desc=f"Generating {binary_name}", 
                 unit='B', unit_scale=True) as pbar:
            binary_data = generate_elf_binary(size_mb)
            with open(f"{output_dir}/{binary_name}", 'wb') as f:
                f.write(binary_data)
            pbar.update(len(binary_data))
        config_lines.append(f"{binary_name}, {size_mb:.2f}MB")
        
        # Generate document files only for appropriate sizes
        if size_mb <= 50:  # Only generate docs for sizes up to 50MB
            for doc_type in ['docx', 'pdf', 'pptx']:
                doc_name = f"document_{size_str}.{doc_type}"
                logging.info(f"  Generating document: {doc_name}")
                doc_content = generate_document(size_mb, doc_type)
                if doc_content:  # Only write if content was generated
                    with open(f"{output_dir}/{doc_name}", 'wb') as f:
                        f.write(doc_content)
                    config_lines.append(f"{doc_name}, {size_mb:.2f}MB")
        
        if size_mb <= 10:  # Only generate Excel for sizes up to 10MB
            doc_name = f"document_{size_str}.xlsx"
            logging.info(f"  Generating document: {doc_name}")
            doc_content = generate_document(size_mb, 'xlsx')
            if doc_content:
                with open(f"{output_dir}/{doc_name}", 'wb') as f:
                    f.write(doc_content)
                config_lines.append(f"{doc_name}, {size_mb:.2f}MB")
        
        # Generate zip file with mixed content
        zip_name = f"archive_{size_str}.zip"
        logging.info(f"  Generating archive: {zip_name}")
        with zipfile.ZipFile(f"{output_dir}/{zip_name}", 'w', zipfile.ZIP_DEFLATED) as zf:
            logging.info(f"    Adding content to {zip_name}")
            if size_mb <= 50:
                # For smaller files, include a mix of content
                split_size = size_mb/4
                zf.writestr('binary1.elf', generate_elf_binary(split_size))
                zf.writestr('binary2.elf', generate_elf_binary(split_size))
                doc_content = generate_document(split_size, 'pdf')
                if doc_content:
                    zf.writestr('document.pdf', doc_content)
                zf.writestr('image.jpg', generate_image(split_size))
            else:
                # For larger files, just split between binaries and images
                split_size = size_mb/3
                zf.writestr('binary1.elf', generate_elf_binary(split_size))
                zf.writestr('binary2.elf', generate_elf_binary(split_size))
                zf.writestr('image.jpg', generate_image(split_size))
        
        config_lines.append(f"{zip_name}, {size_mb:.2f}MB")
        
        logging.info(f"Completed generation for size: {size_str}")
        return config_lines
        
    except Exception as e:
        logging.error(f"Error generating content for size {size_str}: {e}")
        return []

def generate_server_content():
    """Generate various types of content for the server using parallel processing."""
    output_dir = 'server_content'
    os.makedirs(output_dir, exist_ok=True)
    
    sizes = create_size_variants()
    all_config_lines = []
    
    # Use ThreadPoolExecutor for parallel processing
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
        # Submit all tasks
        future_to_size = {
            executor.submit(generate_size_variant, size_kb, output_dir): size_kb 
            for size_kb in sizes
        }
        
        # Collect results as they complete
        for future in tqdm(concurrent.futures.as_completed(future_to_size), total=len(future_to_size), desc="Generating content"):
            size_kb = future_to_size[future]
            try:
                config_lines = future.result()
                all_config_lines.extend(config_lines)
            except Exception as e:
                logging.error(f"Failed to generate content for size {get_size_string(size_kb)}: {e}")

    # Write configuration file
    logging.info("Writing configuration file")
    with open(f'{output_dir}/config.txt', 'w') as config:
        config.write('\n'.join(sorted(all_config_lines)))
    logging.info("Content generation complete")

if __name__ == '__main__':
    generate_server_content()